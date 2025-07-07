# clinica/routes/export.py

# --- Importaciones Necesarias ---
from flask import Blueprint, send_file, request, render_template, current_app
from werkzeug.utils import secure_filename
from io import BytesIO
import pandas as pd
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Importaciones de tus Modelos ---
from ..models import db, Paciente, Evolucion

# --- Creación del Blueprint ---
export_bp = Blueprint('export', __name__)

# --- Exportar a Excel (Tu código original) ---
@export_bp.route('/exportar_excel/<int:id>')
def exportar_excel(id):
    paciente = Paciente.query.get_or_404(id)
    datos = {"Campo": [], "Valor": []}
    campos = [ "id", "nombres", "apellidos", "tipo_documento", "documento", "fecha_nacimiento", "edad", "email", "telefono", "genero", "estado_civil", "direccion", "barrio", "municipio", "departamento", "aseguradora", "tipo_vinculacion", "ocupacion", "referido_por", "nombre_responsable", "telefono_responsable", "parentesco", "motivo_consulta", "enfermedad_actual", "antecedentes_personales", "antecedentes_familiares", "antecedentes_quirurgicos", "antecedentes_hemorragicos", "farmacologicos", "reaccion_medicamentos", "alergias", "habitos", "cepillado", "examen_fisico", "ultima_visita_odontologo", "plan_tratamiento", "observaciones"]
    for campo in campos:
        valor = getattr(paciente, campo, "")
        datos["Campo"].append(campo.replace("_", " ").capitalize())
        datos["Valor"].append(valor if valor else "No disponible")
    df = pd.DataFrame(datos)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Paciente')
    output.seek(0)
    return send_file(output, download_name=f"Paciente_{paciente.id}.xlsx", as_attachment=True)


# --- Exportar a Word (VERSIÓN FINAL CON FORMATO Y LÓGICA DE IMAGEN CORRECTA) ---
@export_bp.route('/exportar_word/<int:id>')
def exportar_word(id):
    paciente = Paciente.query.get_or_404(id)
    doc = Document()

    # Estilo de fuente por defecto
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(8)

    # --- 2. ENCABEZADO CON FECHA, TÍTULO Y SUBTÍTULO ---
    
    # --- ▼▼▼ CAMBIO 1: AÑADIR FECHA DE EMISIÓN ▼▼▼ ---
    fecha_actual = datetime.now().strftime('%d/%m/%Y')
    p_fecha = doc.add_paragraph(f'Fecha de Emisión: {fecha_actual}')
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Hacemos que la fecha tenga un tamaño de fuente más pequeño
    for run in p_fecha.runs:
        run.font.size = Pt(7)
        run.italic = True
    
    # Título Principal
    titulo = doc.add_heading('Historia Clínica Odontológica', level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- ▼▼▼ CAMBIO 2: ESTILO PERSONALIZADO PARA EL CONSULTORIO ▼▼▼ ---
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Añadimos el texto en un 'run' para poder aplicarle estilos específicos
    run_consultorio = subtitulo.add_run('Odontologia Dr. Rueis Pitre')
    # Personalizamos la fuente del consultorio
    font_consultorio = run_consultorio.font
    font_consultorio.name = 'Calibri'  # Puedes cambiar a 'Times New Roman', etc.
    font_consultorio.size = Pt(10)     # Tamaño de fuente 10
    font_consultorio.italic = True
    font_consultorio.bold = True
    
    doc.add_paragraph() # Espacio en blanco

    # Tabla de Datos de Filiación
    doc.add_heading('1. Datos de Filiación', level=2)
    campos_filiacion = [ ("Nombres", paciente.nombres), ("Apellidos", paciente.apellidos), ("Tipo de Documento", paciente.tipo_documento), ("Documento", paciente.documento), ("Fecha de Nacimiento", paciente.fecha_nacimiento), ("Edad", paciente.edad), ("Email", paciente.email), ("Teléfono", paciente.telefono), ("Género", paciente.genero), ("Estado Civil", paciente.estado_civil), ("Dirección", paciente.direccion), ("Ocupación", paciente.ocupacion), ("Aseguradora", paciente.aseguradora), ("Tipo de Vinculación", paciente.tipo_vinculacion) ]
    crear_tabla_formato(doc, campos_filiacion, una_columna=False)
    doc.add_paragraph()

    # Tabla de Anamnesis y Antecedentes
    doc.add_heading('2. Anamnesis y Antecedentes', level=2)
    campos_anamnesis = [ ("Motivo de Consulta", paciente.motivo_consulta), ("Enfermedad Actual", paciente.enfermedad_actual), ("Antecedentes Personales", paciente.antecedentes_personales), ("Antecedentes Familiares", paciente.antecedentes_familiares), ("Antecedentes Quirúrgicos", paciente.antecedentes_quirurgicos), ("Antecedentes Hemorrágicos", paciente.antecedentes_hemorragicos), ("Farmacológicos", paciente.farmacologicos), ("Reacción a Medicamentos", paciente.reaccion_medicamentos), ("Alergias", paciente.alergias), ("Hábitos", paciente.habitos), ("Cepillado", paciente.cepillado), ("Examen Físico", paciente.examen_fisico), ("Última Visita al Odontólogo", paciente.ultima_visita_odontologo), ("Plan de Tratamiento", paciente.plan_tratamiento), ]
    crear_tabla_formato(doc, campos_anamnesis, una_columna=True)
    doc.add_paragraph()

    # --- 3. ANEXOS GRÁFICOS (VERSIÓN MEJORADA CON TABLA) ---
    doc.add_heading('3. Anexos Gráficos', level=2)
    
    # Creamos una tabla de 2 columnas para las imágenes anexas
    tabla_imagenes = doc.add_table(rows=1, cols=2)
    tabla_imagenes.style = 'Table Grid'
    
    # Celda para Imagen 1
    celda_img1 = tabla_imagenes.cell(0, 0)
    add_image_to_cell(celda_img1, paciente.imagen_1, "Imagen Anexa 1", width=3.0)

    # Celda para Imagen 2
    celda_img2 = tabla_imagenes.cell(0, 1)
    add_image_to_cell(celda_img2, paciente.imagen_2, "Imagen Anexa 2", width=3.0)
    
    doc.add_paragraph() # Espacio
    
    # El dentigrama se añade debajo, ocupando todo el ancho
    doc.add_paragraph("Dentigrama:").bold = True
    add_image_to_doc(doc, paciente.dentigrama_canvas, width=6.0) # Ya no necesita etiqueta
    doc.add_paragraph()
    
    
    # Tabla de Evoluciones
    doc.add_heading('4. Evolución del Paciente', level=2)
    tabla_evos = doc.add_table(rows=1, cols=2)
    tabla_evos.style = 'Table Grid'
    tabla_evos.columns[0].width = Inches(1.25)
    tabla_evos.columns[1].width = Inches(5.25)
    hdr_cells = tabla_evos.rows[0].cells
    hdr_cells[0].text = 'Fecha'; hdr_cells[0].paragraphs[0].runs[0].bold = True
    hdr_cells[1].text = 'Descripción de la Evolución'; hdr_cells[1].paragraphs[0].runs[0].bold = True
    for evo in paciente.evoluciones.order_by(Evolucion.fecha.asc()):
        row_cells = tabla_evos.add_row().cells
        row_cells[0].text = evo.fecha.strftime('%d/%m/%Y %H:%M')
        row_cells[1].text = evo.descripcion

    # Generación del archivo
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    download_filename = f"Historia_Clinica_{paciente.documento or paciente.id}.docx"
    return send_file(output, download_name=download_filename, as_attachment=True)


# --- FUNCIONES AUXILIARES (DEFINITIVAS) ---

def crear_tabla_formato(doc, campos, una_columna=False):
    """Crea una tabla con etiquetas en negrita y valores para los campos proporcionados."""
    cols = 2 if una_columna else 4
    tabla = doc.add_table(rows=0, cols=cols)
    tabla.style = 'Table Grid'
    
    paso = 1 if una_columna else 2
    for i in range(0, len(campos), paso):
        row_cells = tabla.add_row().cells
        
        # Campo de la izquierda
        label_izq, value_izq = campos[i]
        p_label_izq = row_cells[0].paragraphs[0]
        p_label_izq.add_run(f"{label_izq}:").bold = True
        row_cells[1].text = str(value_izq) if value_izq is not None else 'N/A'
        
        # Campo de la derecha (si aplica)
        if not una_columna and i + 1 < len(campos):
            label_der, value_der = campos[i + 1]
            p_label_der = row_cells[2].paragraphs[0]
            p_label_der.add_run(f"{label_der}:").bold = True
            row_cells[3].text = str(value_der) if value_der is not None else 'N/A'

# --- FUNCIÓN AUXILIAR MODIFICADA ---
# Reemplaza tu add_image_to_doc por esta nueva versión
def add_image_to_doc(doc, ruta_relativa_db, width=3.0):
    """Añade una imagen directamente al documento."""
    if not ruta_relativa_db:
        doc.add_paragraph("(No disponible)")
        return
    ruta_absoluta = os.path.join(current_app.root_path, 'static', ruta_relativa_db)
    if os.path.exists(ruta_absoluta):
        try:
            doc.add_picture(ruta_absoluta, width=Inches(width))
        except Exception as e:
            doc.add_paragraph(f"(Error: {e})")
    else:
        doc.add_paragraph("(Archivo no encontrado)")

# --- NUEVA FUNCIÓN AUXILIAR ---
# Añade esta nueva función a tu archivo
def add_image_to_cell(cell, ruta_relativa_db, label, width=3.0):
    """Añade una etiqueta y una imagen dentro de una celda de tabla."""
    # Limpiamos el párrafo por defecto de la celda
    cell.text = ''
    p = cell.add_paragraph()
    p.add_run(f"{label}:").bold = True
    
    if not ruta_relativa_db:
        cell.add_paragraph("(No disponible)")
        return
    
    ruta_absoluta = os.path.join(current_app.root_path, 'static', ruta_relativa_db)

    if os.path.exists(ruta_absoluta):
        try:
            # Añadimos la imagen en un nuevo párrafo para mejor control
            cell.add_paragraph().add_run().add_picture(ruta_absoluta, width=Inches(width))
        except Exception as e:
            cell.add_paragraph(f"(Error: {e})")
    else:
        cell.add_paragraph("(Archivo no encontrado)")